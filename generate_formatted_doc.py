"""
Generate SRDC Documentation in IQAC Format Template Structure.
This script produces a properly formatted .docx file matching the instructor's template.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ============================================================
# STYLE CONFIGURATION
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# Configure heading styles
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.name = 'Times New Roman'
    heading_style.font.color.rgb = RGBColor(0, 0, 0)
    heading_style.font.bold = True
    if i == 1:
        heading_style.font.size = Pt(16)
    elif i == 2:
        heading_style.font.size = Pt(14)
    else:
        heading_style.font.size = Pt(12)

# Set narrow margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def add_centered_text(text, size=12, bold=False, space_after=0, space_before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = 'Times New Roman'
    return p

def add_normal_text(text, size=12, bold=False, space_after=6, space_before=0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = 'Times New Roman'
    return p

def add_heading_custom(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_page_break():
    doc.add_page_break()

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

# ============================================================
# PAGE 1: TITLE PAGE
# ============================================================
doc.add_paragraph()  # spacing
add_centered_text("A Real-time Research Project/Societal Related Project", 14, True, space_after=6)
add_centered_text("Report on", 14, False, space_after=12)
add_centered_text("SRDC: SEMANTICS-BASED RANSOMWARE DETECTION", 16, True, space_after=2)
add_centered_text("AND CLASSIFICATION WITH LLM-ASSISTED PRE-TRAINING", 16, True, space_after=18)
add_centered_text("Submitted in Partial fulfillment of requirements for B.Tech II Year II Semester course", 12, False, space_after=24)

add_centered_text("By", 12, True, space_after=12)

students = [
    ("1. AELUGU RANITH KUMAR", "25BD5A6615"),
    ("2. CHINTALA RISHITH", "25BD5A6617"),
    ("3. G SAMPATH", "25BD5A6618"),
    ("4. KONDURI ABHIRAM", "25BD5A6620"),
    ("5. MAINENI SREENILAY", "25BD5A6621"),
]
for name, roll in students:
    add_centered_text(f"{name} ({roll})", 12, False, space_after=2)

doc.add_paragraph()
add_centered_text("Under the guidance of", 12, False, space_after=6, space_before=12)
add_centered_text("Mr. Badrinath", 12, True, space_after=2)
add_centered_text("Assistant Professor, Department of CSE", 12, False, space_after=24)

doc.add_paragraph()
doc.add_paragraph()

add_centered_text("KESHAV MEMORIAL INSTITUTE OF TECHNOLOGY", 12, True, space_after=2)
add_centered_text("(AN AUTONOMOUS INSTITUTION)", 11, True, space_after=2)
add_centered_text("Accredited by NBA & NAAC, Approved by AICTE, Affiliated to JNTUH.", 10, False, space_after=2)
add_centered_text("Narayanaguda, Hyderabad, Telangana-29", 11, False, space_after=2)
add_centered_text("2025-26", 12, True, space_after=0)

add_page_break()

# ============================================================
# PAGE 2: CERTIFICATE
# ============================================================
add_centered_text("KESHAV MEMORIAL INSTITUTE OF TECHNOLOGY", 12, True, space_after=2, space_before=12)
add_centered_text("(AN AUTONOMOUS INSTITUTION)", 11, True, space_after=2)
add_centered_text("Accredited by NBA & NAAC, Approved by AICTE, Affiliated to JNTUH", 10, False, space_after=2)
add_centered_text("Narayanaguda, Hyderabad, Telangana-29", 11, False, space_after=24)

add_centered_text("CERTIFICATE", 16, True, space_after=18)

add_normal_text(
    'This is to certify that this is a bonafide record of the project report titled '
    '"SRDC: Semantics-Based Ransomware Detection and Classification with LLM-Assisted Pre-Training" '
    'which is being presented as the Real-time Research Project / Societal Related Project report by',
    12, False, space_after=12
)

for name, roll in students:
    add_centered_text(f"{name} ({roll})", 12, False, space_after=2)

doc.add_paragraph()
add_normal_text(
    "In partial fulfillment for the II Year II Semester Course RTRP in KMIT affiliated to the "
    "Jawaharlal Nehru Technological University, Hyderabad",
    12, False, space_after=24
)

doc.add_paragraph()
# Mentor and coordinator line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Mentor")
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("(Mr. Badrinath)")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run2 = p.add_run("                                                              Program Coordinator")
run2.font.size = Pt(12)
run2.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("(Mr. Shailesh Gangakhedkar)")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()
add_normal_text("Submitted for Final Project Review held on ____________________________", 12, False, space_after=0)

add_page_break()

# ============================================================
# PAGE 3: VISION & MISSION OF KMIT  (WAS MISSING)
# ============================================================
add_centered_text("Vision & Mission of KMIT", 16, True, space_after=18)

add_normal_text("Vision", 14, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "To be the fountain head in producing highly skilled, globally competent engineers.",
    12, False, space_after=6
)
add_normal_text(
    "Producing quality graduates trained in the latest software technologies and related tools "
    "and striving to make India a world leader in software products and services.",
    12, False, space_after=12
)

add_normal_text("Mission", 14, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)

mission_points = [
    "To provide a learning environment that inculcates problem solving skills, professional, ethical responsibilities, lifelong learning through multi model platforms and prepares students to become successful professionals.",
    "To establish an industry institute Interaction to make students ready for the industry.",
    "To provide exposure to students on the latest hardware and software tools.",
    "To promote research-based projects/activities in the emerging areas of technology convergence.",
    "To encourage and enable students to not merely seek jobs from the industry but also to create new enterprises.",
    "To induce a spirit of nationalism which will enable the student to develop, understand India's challenges and to encourage them to develop effective solutions.",
    "To support the faculty to accelerate their learning curve to deliver excellent service to students.",
]
for m in mission_points:
    p = doc.add_paragraph(m, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

add_page_break()

# ============================================================
# PAGE 4: PROGRAM OUTCOMES (FULL TEXT)
# ============================================================
add_centered_text("PROGRAM OUTCOMES (POs)", 16, True, space_after=18)

program_outcomes = [
    ("PO1. Engineering Knowledge:", "Apply the knowledge of mathematics, science, engineering fundamentals, and an engineering specialization to the solution of complex engineering problems."),
    ("PO2. Problem Analysis:", "Identify, formulate, review research literature, and analyze complex engineering problems reaching substantiated conclusions using first principles of mathematics, natural sciences, and engineering sciences."),
    ("PO3. Design/Development of Solutions:", "Design solutions for complex engineering problems and design system components or processes that meet the specified needs with appropriate consideration for the public health and safety, and the cultural, societal, and environmental considerations."),
    ("PO4. Conduct Investigations of Complex Problems:", "Use research-based knowledge and research methods including design of experiments, analysis and interpretation of data, and synthesis of the information to provide valid conclusions."),
    ("PO5. Modern Tool Usage:", "Create, select, and apply appropriate techniques, resources, and modern engineering and IT tools including prediction and modeling to complex engineering activities with an understanding of the limitations."),
    ("PO6. The Engineer and Society:", "Apply reasoning informed by contextual knowledge to assess societal, health, safety, legal and cultural issues and the consequent responsibilities relevant to professional engineering practice."),
    ("PO7. Environment and Sustainability:", "Understand the impact of the professional engineering solutions in societal and environmental contexts and demonstrate the knowledge of, and need for sustainable development."),
    ("PO8. Ethics:", "Apply ethical principles and commit to professional ethics and responsibilities and norms of the engineering practice."),
    ("PO9. Individual and Team Work:", "Function effectively as an individual, and as a member or leader in diverse teams and in multidisciplinary settings."),
    ("PO10. Communication:", "Communicate effectively on complex engineering activities with the engineering community and with society at large, such as, being able to comprehend and write effective reports and design documentation, make effective presentations, and give and receive clear instructions."),
    ("PO11. Project Management and Finance:", "Demonstrate knowledge and understanding of the engineering and management principles and apply these to one's own work, as a member and leader in a team, to manage projects and in multidisciplinary environments."),
    ("PO12. Life-Long Learning:", "Recognize the need for, and have the preparation and ability to engage in independent and life-long learning in the broadest context of technological change."),
]

for title, desc in program_outcomes:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    run_title = p.add_run(title + " ")
    run_title.bold = True
    run_title.font.size = Pt(12)
    run_title.font.name = 'Times New Roman'
    run_desc = p.add_run(desc)
    run_desc.font.size = Pt(12)
    run_desc.font.name = 'Times New Roman'

add_page_break()

# ============================================================
# PAGE 5: PROJECT OUTCOMES
# ============================================================
add_centered_text("PROJECT OUTCOMES", 16, True, space_after=18)

project_outcomes = [
    "P1: Apply GPT-2 based language models for semantic analysis of dynamic ransomware behaviors.",
    "P2: Implement internal semantic processing to convert raw binary features into natural language representations.",
    "P3: Design and execute zero-day ransomware detection experiments with held-out ransomware families.",
    "P4: Build an interactive sandbox simulation demo for real-time ransomware detection and family classification.",
]
for po in project_outcomes:
    add_normal_text(po, 12, False, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)

doc.add_paragraph()

# ============================================================
# PO MAPPING TABLE
# ============================================================
add_centered_text("MAPPING PROJECT OUTCOMES WITH PROGRAM OUTCOMES", 14, True, space_after=12)

# Mapping data: P1-P4 x PO1-PO12
mapping = {
    'P1': ['H', 'H', 'M', 'H', 'H', 'L', '', '', 'M', '', '', 'M'],
    'P2': ['H', 'H', 'H', 'M', 'H', '', '', '', '', '', '', 'M'],
    'P3': ['M', 'H', 'H', 'H', 'H', '', '', '', 'M', '', '', 'H'],
    'P4': ['M', 'M', 'H', 'M', 'H', '', '', '', 'H', 'H', 'M', ''],
}

po_headers = ['PO'] + [f'PO{i}' for i in range(1, 13)]
table = doc.add_table(rows=5, cols=13, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
for i, header in enumerate(po_headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
    set_cell_shading(cell, 'D9E2F3')

# Data rows
for row_idx, (pname, values) in enumerate(mapping.items(), start=1):
    table.rows[row_idx].cells[0].text = pname
    for p in table.rows[row_idx].cells[0].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
    for col_idx, val in enumerate(values, start=1):
        cell = table.rows[row_idx].cells[col_idx]
        cell.text = val
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'

doc.add_paragraph()
add_centered_text("L – LOW     M – MEDIUM     H – HIGH", 10, True, space_after=0)

add_page_break()

# ============================================================
# DECLARATION
# ============================================================
add_centered_text("DECLARATION", 16, True, space_after=18)

add_normal_text(
    'We hereby declare that the results embodied in the dissertation entitled '
    '"SRDC: Semantics-Based Ransomware Detection and Classification with LLM-Assisted Pre-Training" '
    'has been carried out by us together during the academic year 2024-25 as a partial fulfillment of the '
    'B.Tech II Year II Semester Course "Real-time Research Project / Societal Related Project". '
    'We have not submitted this report to any other Course/College.',
    12, False, space_after=24
)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Student Name                                          Roll No.")
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()
for i, (name, roll) in enumerate(students, 1):
    add_normal_text(f"{name} ({roll})", 12, False, space_after=4, alignment=WD_ALIGN_PARAGRAPH.LEFT)

add_page_break()

# ============================================================
# ACKNOWLEDGEMENT
# ============================================================
add_centered_text("ACKNOWLEDGEMENT", 16, True, space_after=18)

ack_paragraphs = [
    "We take this opportunity to thank all the people who have rendered their full support to our project work. We render our thanks to Dr. B L Malleswari, Principal who encouraged us to do the Project.",
    "We are grateful to Mr. Neil Gogte, Founder & Director and Mr. S. Nitin, Director, for facilitating all the amenities required for carrying out this project.",
    "We express our sincere gratitude to Ms. Deepa Ganu, Academic Director for providing an excellent environment in the college.",
    "We are also thankful to Mr. Shailesh Gangakhedkar, Real-Time Research Project Program Coordinator for providing us with time to make this project a success within the given schedule.",
    "We are also thankful to our Project Mentor Mr. Badrinath, for his valuable guidance and encouragement given to us throughout the project work.",
    "We would like to thank the entire KMIT faculty, who helped us directly and indirectly in the completion of the project.",
    "We sincerely thank our friends and family for their constant motivation during the project work.",
]
for para in ack_paragraphs:
    add_normal_text(para, 12, False, space_after=10)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Student Name                                          Roll No.")
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()
for name, roll in students:
    add_normal_text(f"{name} ({roll})", 12, False, space_after=4, alignment=WD_ALIGN_PARAGRAPH.LEFT)

add_page_break()

# ============================================================
# ABSTRACT
# ============================================================
add_centered_text("ABSTRACT", 16, True, space_after=18)

abstract_paragraphs = [
    'In this project, we implement and reproduce the research paper "SRDC: Semantics-based Ransomware Detection and Classification with LLM-assisted Pre-training" published at AAAI 2025. The system addresses the critical cybersecurity challenge of detecting zero-day ransomware that evades traditional signature-based antivirus solutions.',
    "The proposed approach uses Internal Semantic Processing to convert raw dynamic behavior logs (API calls, registry modifications, file operations, dropped files) into natural language sentences. These semantically enriched descriptions are then fed into a GPT-2 model (zhouce/RDC-GPT) that has been pre-trained on cybersecurity domain knowledge from Microsoft documentation and expert-verified corpus.",
    "We fine-tuned two separate classifiers: (1) a Binary Detection model that classifies program behavior as Ransomware or Goodware, achieving 97% accuracy even on unseen zero-day ransomware families, and (2) a Family Classification model that categorizes detected ransomware into one of 12 known families (e.g., CryptoWall, TeslaCrypt, Reveton, CryptLocker) with 60.7% balanced accuracy.",
    "The system is demonstrated through an interactive sandbox simulation that processes live behavior samples and provides real-time detection with confidence scores. The project validates that LLM-based semantic understanding of program behavior significantly outperforms traditional machine learning approaches for ransomware detection.",
    'To transition this research into a deployable cybersecurity application, the core AI engine was productized during Phase 1 into a client-server system named "SRDC Shield". This architecture combines the fine-tuned deep learning models with a real-time Google Chrome Manifest V3 extension frontend and a local Python Flask backend service. The combined application actively intercepts high-risk web files before they execute, proving that semantic LLM models can move beyond offline validation to provide proactive perimeter protection.',
]
for para in abstract_paragraphs:
    add_normal_text(para, 12, False, space_after=8)

add_page_break()

# ============================================================
# LIST OF FIGURES
# ============================================================
add_centered_text("LIST OF FIGURES", 16, True, space_after=18)

figures_table = doc.add_table(rows=6, cols=3, style='Table Grid')
figures_table.alignment = WD_TABLE_ALIGNMENT.CENTER

fig_headers = ['S.No', 'Name of Screenshot', 'Page No.']
for i, h in enumerate(fig_headers):
    cell = figures_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
    set_cell_shading(cell, 'D9E2F3')

fig_names = [
    "System Architecture Diagram",
    "Internal Semantic Processing Flow",
    "Training Loss Curve",
    "Demo Output - Goodware Detection",
    "Demo Output - Ransomware Detection",
]
for i, fname in enumerate(fig_names, 1):
    figures_table.rows[i].cells[0].text = str(i)
    figures_table.rows[i].cells[1].text = fname
    figures_table.rows[i].cells[2].text = ""
    for j in range(3):
        for p in figures_table.rows[i].cells[j].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'

add_page_break()

# ============================================================
# TABLE OF CONTENTS (Restructured to match template)
# ============================================================
add_centered_text("CONTENTS", 16, True, space_after=18)

toc_entries = [
    ("CHAPTER – 1", True, False),
    ("INTRODUCTION", True, False),
    ("    Purpose of the Project", False, False),
    ("    Problem with Existing Systems", False, False),
    ("    Proposed System", False, False),
    ("    Scope of the Project", False, False),
    ("    Architecture Diagram", False, False),
    ("CHAPTER – 2", True, False),
    ("LITERATURE SURVEY", True, False),
    ("CHAPTER – 3", True, False),
    ("SOFTWARE REQUIREMENT SPECIFICATION", True, False),
    ("    Introduction to SRS", False, False),
    ("    Role of SRS", False, False),
    ("    Requirements Specification Document", False, False),
    ("    Functional Requirements", False, False),
    ("    Non-Functional Requirements", False, False),
    ("    Performance Requirements", False, False),
    ("    Software Requirements", False, False),
    ("    Hardware Requirements", False, False),
    ("CHAPTER – 4", True, False),
    ("SYSTEM DESIGN", True, False),
    ("    Introduction to UML", False, False),
    ("    UML Diagrams", False, False),
    ("    Use Case Diagram", False, False),
    ("    Sequence Diagram", False, False),
    ("    State Chart Diagram", False, False),
    ("    System Architecture", False, False),
    ("    TECHNOLOGIES USED", True, False),
    ("CHAPTER – 5", True, False),
    ("IMPLEMENTATION", True, False),
    ("    Environment Setup", False, False),
    ("    Dataset Acquisition", False, False),
    ("    Internal Semantic Processing", False, False),
    ("    Data Splitting", False, False),
    ("    Model Training", False, False),
    ("    Interactive Demo", False, False),
    ("    SRDC Shield - Production Extension", False, False),
    ("    Problems Faced & Solutions", False, False),
    ("    Screenshots", False, False),
    ("CHAPTER – 6", True, False),
    ("SOFTWARE TESTING", True, False),
    ("    Introduction", False, False),
    ("    Testing Objectives", False, False),
    ("    Testing Strategies", False, False),
    ("    System Evaluation", False, False),
    ("    Test Cases", False, False),
    ("CONCLUSION", True, False),
    ("FUTURE ENHANCEMENTS", True, False),
    ("REFERENCES", True, False),
    ("BIBLIOGRAPHY", True, False),
]

for entry_text, is_bold, _ in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(entry_text)
    run.bold = is_bold
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

add_page_break()

# ============================================================
# CHAPTER 1: INTRODUCTION
# ============================================================
add_centered_text("CHAPTER – 1", 16, True, space_after=6)
add_centered_text("INTRODUCTION", 16, True, space_after=18)

add_normal_text(
    "Ransomware is one of the most devastating forms of malware in the modern cybersecurity landscape. "
    "It encrypts victims' files and demands a ransom payment for the decryption key. Traditional antivirus "
    "solutions rely on static signature matching - comparing suspicious files against a database of known "
    "malware signatures. While effective against known threats, this approach completely fails against zero-day "
    "ransomware: new, previously unseen variants that have no existing signature in any database.",
    12, False, space_after=8
)
add_normal_text(
    "The SRDC (Semantics-based Ransomware Detection and Classification) system addresses this fundamental "
    "limitation by shifting the detection paradigm from static signatures to behavioral semantics. Instead of "
    'asking "Does this code match a known virus?", the system asks "What is this program trying to do, and '
    'does its behavior look malicious?"',
    12, False, space_after=8
)
add_normal_text(
    "The core innovation lies in using a Large Language Model (LLM) - specifically a fine-tuned GPT-2 "
    "architecture - to semantically understand program behavior. Raw dynamic behavior logs (API calls, "
    "registry modifications, file operations, and dropped files) are first translated into natural English "
    "sentences through Internal Semantic Processing. These sentences are then analyzed by the GPT-2 model, "
    "which has been pre-trained on cybersecurity domain knowledge, enabling it to understand malicious intent "
    "the same way a human security expert would.",
    12, False, space_after=8
)
add_normal_text(
    'This project implements and reproduces the research paper "SRDC: Semantics-based Ransomware Detection '
    'and Classification with LLM-assisted Pre-training" published at AAAI 2025 (Ce Zhou et al.). The '
    'implementation covers the complete pipeline from raw data processing to a real-time interactive detection demo.',
    12, False, space_after=12
)

# 1.1 Purpose of the Project
add_normal_text("1.1 Purpose of the Project", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "The purpose of this project is to implement and validate a semantics-based ransomware detection system "
    "that uses LLM-assisted pre-training to detect zero-day ransomware with high accuracy. The system aims to "
    "demonstrate that semantic understanding of program behavior using GPT-2 significantly outperforms "
    "traditional signature-based and statistical machine learning approaches.",
    12, False, space_after=8
)

# 1.2 Problem with Existing Systems
add_normal_text("1.2 Problem with Existing Systems", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "Traditional signature-based ransomware detection fails against zero-day attacks because it can only "
    "identify threats it has previously seen. There is a critical need for an intelligent detection system "
    "that can identify malicious intent from program behavior, even when the specific ransomware family "
    "has never been encountered before. Existing machine learning approaches using Random Forest and SVM "
    "classifiers struggle with zero-day detection due to their reliance on statistical patterns rather "
    "than semantic understanding.",
    12, False, space_after=8
)

# 1.3 Proposed System
add_normal_text("1.3 Proposed System", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "The proposed SRDC system uses Internal Semantic Processing to convert raw dynamic behavior logs "
    "(API calls, registry modifications, file operations, dropped files) into natural language sentences. "
    "These semantically enriched descriptions are fed into a GPT-2 model (zhouce/RDC-GPT) pre-trained on "
    "cybersecurity domain knowledge. The system performs dual classification: (1) Binary Detection "
    "(Ransomware vs Goodware) achieving 97% accuracy on zero-day samples, and (2) Family Classification "
    "across 12 known ransomware families.",
    12, False, space_after=8
)

# 1.4 Scope of the Project
add_normal_text("1.4 Scope of the Project", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "The scope includes implementing the complete SRDC pipeline: dataset preprocessing with internal "
    "semantic processing, data splitting for zero-day simulation, fine-tuning the GPT-2 model for binary "
    "detection and family classification, and building an interactive sandbox demo for real-time ransomware detection.",
    12, False, space_after=8
)
add_normal_text(
    "The scope has been extended beyond offline pipeline simulation to incorporate real-time endpoint "
    "protection mechanics. This includes the development of a production-ready Manifest V3 browser extension "
    "for proactive traffic interception, a local Flask backend handling multi-tier security evaluations "
    "(Signature scanning via the VirusTotal API combined with Zero-Day Semantic AI scanning), a secure "
    "workspace for nested archive parsing, and data processing routines to maintain system performance on local devices.",
    12, False, space_after=8
)

# 1.5 Architecture Diagram
add_normal_text("1.5 Architecture Diagram", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "The SRDC system architecture consists of three main layers:",
    12, False, space_after=6
)
add_normal_text(
    "1. Data Layer: Raw CSV dataset (RansomwareData.csv) containing 1,524 samples with 30,000+ binary "
    "features representing API calls, registry operations, file operations, dropped files, file extensions, "
    "directories, and embedded strings.",
    12, False, space_after=4
)
add_normal_text(
    "2. Processing Layer: Internal Semantic Processing module that converts binary feature vectors into "
    "natural language sentences. The split_data module creates stratified train/test splits with zero-day "
    "family hold-out.",
    12, False, space_after=4
)
add_normal_text(
    "3. Model Layer: Two fine-tuned GPT-2 classifiers - Binary Detection (2-class: Goodware/Ransomware) "
    "and Family Classification (12-class). Both use mean pooling of the last hidden states followed by a "
    "linear classification head.",
    12, False, space_after=8
)
add_normal_text(
    'The architecture was evolved into an integrated production application called "SRDC Shield" with a '
    'Dual-Layer Client-Server framework:',
    12, False, space_after=6
)
add_normal_text(
    "1. Client Tier (Chrome Extension Frontend): Built using a modern, glassmorphic HTML/CSS layout, this "
    "component runs a background service worker that hooks into the browser's download streams. It intercepts "
    'target file formats (.exe, .msi, .zip) pre-disk, instantly forcing them into a paused, unexecutable '
    '".crdownload" state to insulate the system while a scan takes place.',
    12, False, space_after=4
)
add_normal_text(
    "2. Production Server Tier (Flask Backend Engine): A high-performance Python application gateway that "
    "preloads both fine-tuned model checkpoints (srdc_zero_day_BEST.pth and srdc_family_BEST.pth) straight "
    "into system RAM on initialization to eliminate cold-start lag. It maps file metadata across a Dual-Layer "
    "Security routine: Layer 1 runs a fast SHA-256 hash comparison against the live VirusTotal registry API, "
    "while Layer 2 orchestrates structural feature parsing via the pefile library, converts the results into "
    "text descriptions, and passes them to the GPT-2 transformer network for live classification.",
    12, False, space_after=8
)

add_page_break()

# ============================================================
# CHAPTER 2: LITERATURE SURVEY
# ============================================================
add_centered_text("CHAPTER – 2", 16, True, space_after=6)
add_centered_text("LITERATURE SURVEY", 16, True, space_after=18)

add_normal_text(
    "Ransomware attacks have caused billions of dollars in damages worldwide. The WannaCry attack of 2017 "
    "alone affected over 200,000 computers across 150 countries. Traditional detection methods based on "
    "static signatures are reactive by nature - they can only detect threats after a signature has been created, "
    "leaving a critical window of vulnerability during zero-day attacks. The fundamental problem is: How can "
    "we detect ransomware that has never been seen before, based solely on what the program is trying to do?",
    12, False, space_after=12
)

add_normal_text(
    'Sgandurra et al. (2016) - EldeRan: Created the foundational dataset of 1,524 dynamic behavior samples '
    '(582 ransomware from 11 families + 942 goodware) by running executables in a controlled sandbox and '
    'recording API calls, registry modifications, file operations, and dropped files. This dataset forms the '
    'basis of our implementation.',
    12, False, space_after=8
)
add_normal_text(
    "Chen et al. (2018) - Used Random Forest and SVM classifiers on behavioral features for ransomware "
    "detection. While achieving reasonable accuracy on known families, these models struggled with zero-day "
    "detection due to their reliance on statistical patterns rather than semantic understanding.",
    12, False, space_after=8
)
add_normal_text(
    "Ce Zhou et al. (2025) - SRDC Paper (AAAI 2025): Proposed using GPT-2 with domain-specific pre-training "
    "on cybersecurity knowledge corpus. The key innovations include: (1) Internal Semantic Processing to "
    "convert binary features to natural language, (2) External knowledge pre-training using Microsoft "
    "documentation and expert-verified descriptions, and (3) A dual-task framework for both zero-day "
    "detection and family classification.",
    12, False, space_after=12
)

add_normal_text("Base Paper", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    'Title: SRDC: Semantics-based Ransomware Detection and Classification with LLM-assisted Pre-training\n'
    'Authors: Ce Zhou et al.\n'
    'Conference: AAAI 2025 (Association for the Advancement of Artificial Intelligence)\n'
    'Key Contribution: Demonstrated that LLM-based semantic understanding of program behavior significantly '
    'outperforms traditional ML approaches for zero-day ransomware detection, achieving state-of-the-art results.',
    12, False, space_after=8
)

add_page_break()

# ============================================================
# CHAPTER 3: SOFTWARE REQUIREMENT SPECIFICATION (NEW SRS FRAMING)
# ============================================================
add_centered_text("CHAPTER – 3", 16, True, space_after=6)
add_centered_text("SOFTWARE REQUIREMENT SPECIFICATION", 16, True, space_after=18)

# 3.1 Introduction to SRS
add_normal_text("3.1 Introduction to SRS", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "A Software Requirements Specification (SRS) is a document that describes the intended purpose, "
    "requirements, and nature of a software system to be developed. It provides a complete description "
    "of the behavior of the system. The SRS for the SRDC project defines all functional and non-functional "
    "requirements necessary for building a semantics-based ransomware detection and classification system "
    "using LLM-assisted pre-training.",
    12, False, space_after=8
)

# 3.2 Role of SRS
add_normal_text("3.2 Role of SRS", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "The SRS serves as a contract between the development team and the stakeholders. It reduces the development "
    "effort by providing a clear understanding of the system's requirements. It serves as the basis for "
    "estimating costs and schedules, provides a baseline for validation and verification, and facilitates "
    "the transfer of the software to new users or machines. For the SRDC project, the SRS ensures that the "
    "implementation faithfully reproduces the research paper's pipeline while meeting real-time performance "
    "requirements for the production browser extension.",
    12, False, space_after=8
)

# 3.3 Requirements Specification Document
add_normal_text("3.3 Requirements Specification Document", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "The system shall implement the complete SRDC pipeline as described in the AAAI 2025 research paper by "
    "Ce Zhou et al. The implementation covers data preprocessing, internal semantic processing, model training, "
    "inference, and a real-time interactive demo. The production extension adds browser-level file interception "
    "and multi-tier scanning capabilities.",
    12, False, space_after=8
)

# 3.4 Functional Requirements
add_normal_text("3.4 Functional Requirements", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
func_reqs = [
    "FR1: The system shall convert raw binary feature vectors (API calls, registry keys, file operations, dropped files) into natural language sentences using Internal Semantic Processing.",
    "FR2: The system shall perform binary classification to detect whether a given behavior sample is Ransomware or Goodware.",
    "FR3: The system shall classify detected ransomware into one of 12 known families (CryptoWall, TeslaCrypt, Reveton, CryptLocker, etc.).",
    "FR4: The system shall provide confidence scores for all predictions.",
    "FR5: The system shall support zero-day detection by identifying ransomware families never seen during training.",
    "FR6: The browser extension shall intercept .exe, .msi, and .zip downloads and pause them before disk write for scanning.",
    "FR7: The backend shall perform dual-layer scanning: VirusTotal signature check (Layer 1) and GPT-2 semantic analysis (Layer 2).",
    "FR8: The system shall provide an interactive sandbox simulation for demonstration purposes.",
]
for fr in func_reqs:
    p = doc.add_paragraph(fr, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

doc.add_paragraph()

# 3.5 Non-Functional Requirements
add_normal_text("3.5 Non-Functional Requirements", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
nfr = [
    "NFR1: The system shall achieve at least 95% accuracy on zero-day ransomware detection.",
    "NFR2: The backend inference shall complete within 2 seconds per sample on a standard CPU.",
    "NFR3: The browser extension shall not degrade normal browsing performance by more than 5%.",
    "NFR4: The system shall handle binary files up to 100MB without crashing.",
    "NFR5: All temporary extracted files shall be securely deleted after analysis.",
]
for n in nfr:
    p = doc.add_paragraph(n, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

doc.add_paragraph()

# 3.6 Performance Requirements
add_normal_text("3.6 Performance Requirements", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
perf_reqs = [
    "Binary detection accuracy: 97% on zero-day samples (matching research paper benchmarks).",
    "Family classification balanced accuracy: 60.7% across 12 imbalanced classes.",
    "Feature extraction latency: Under 0.05 seconds per file (using 256KB read optimization).",
    "Model loading time: Under 10 seconds on cold start (models preloaded into RAM).",
    "Concurrent scan support: At least 3 simultaneous file scans without performance degradation.",
]
for pr in perf_reqs:
    p = doc.add_paragraph(pr, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

doc.add_paragraph()

# 3.7 Software Requirements
add_normal_text("3.7 Software Requirements", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
sw_reqs = [
    "Python 3.12+",
    "PyTorch (torch) – Deep learning framework",
    "HuggingFace Transformers – Pre-trained GPT-2 model loading",
    "Pandas – Data manipulation and CSV processing",
    "scikit-learn – Data splitting and evaluation metrics",
    "NumPy – Numerical computations",
    "Matplotlib – Visualization",
    "Flask – Backend API server",
    "pefile – PE binary parsing library",
    "Google Colab (for GPU training)",
    "Google Chrome (for Manifest V3 extension)",
]
for s in sw_reqs:
    p = doc.add_paragraph(s, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

doc.add_paragraph()

# 3.8 Hardware Requirements
add_normal_text("3.8 Hardware Requirements", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
hw_reqs = [
    "Minimum 8GB RAM (16GB recommended)",
    "Intel i5 or above processor",
    "Internet connection (for downloading pre-trained models from HuggingFace)",
    "GPU (NVIDIA T4 or above) for model training (Google Colab used)",
    "2GB free disk space for model weights (.pth files)",
]
for h in hw_reqs:
    p = doc.add_paragraph(h, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

add_page_break()

# ============================================================
# CHAPTER 4: SYSTEM DESIGN
# ============================================================
add_centered_text("CHAPTER – 4", 16, True, space_after=6)
add_centered_text("SYSTEM DESIGN", 16, True, space_after=18)

# 4.1 Introduction to UML
add_normal_text("4.1 Introduction to UML", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "Unified Modeling Language (UML) is a standardized modeling language used to visualize the design of a "
    "system. It helps in understanding system architecture and workflows clearly before implementation. "
    "The following UML diagrams describe the SRDC system's structure and behavior.",
    12, False, space_after=12
)

# 4.2 UML Diagrams
add_normal_text("4.2 UML Diagrams", 13, True, space_after=8, alignment=WD_ALIGN_PARAGRAPH.LEFT)

# 4.2.1 Use Case Diagram
add_normal_text("4.2.1 Use Case Diagram", 12, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text("Actors:", 12, True, space_after=4, alignment=WD_ALIGN_PARAGRAPH.LEFT)
for actor in ["Security Analyst (Primary User)", "SRDC System (Automated)"]:
    p = doc.add_paragraph(actor, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

add_normal_text("Use Cases:", 12, True, space_after=4, alignment=WD_ALIGN_PARAGRAPH.LEFT)
use_cases = [
    "Upload/Input behavior sample",
    "Run binary detection (Ransomware vs Goodware)",
    "Run family classification",
    "View detection results with confidence scores",
    "View recommended actions",
]
for uc in use_cases:
    p = doc.add_paragraph(uc, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

add_normal_text(
    "Explanation: The security analyst submits a suspicious behavior sample to the sandbox. The SRDC system "
    "processes it through the GPT-2 pipeline and returns detection results with confidence scores and "
    "recommended actions (e.g., ISOLATE SYSTEM).",
    12, False, space_after=10
)

# 4.2.2 Sequence Diagram
add_normal_text("4.2.2 Sequence Diagram", 12, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text("Steps:", 12, True, space_after=4, alignment=WD_ALIGN_PARAGRAPH.LEFT)
seq_steps = [
    "User submits behavior sample to sandbox",
    "System extracts API, registry, file, and extension features",
    "Internal Semantic Processing converts features to natural language",
    "Text is tokenized using GPT-2 tokenizer (max 1024 tokens)",
    "Binary Detection Model classifies: Ransomware or Goodware",
    "If Ransomware: Family Classification Model identifies the family",
    "Results with confidence scores are returned to the user",
]
for i, step in enumerate(seq_steps, 1):
    add_normal_text(f"  {i}. {step}", 12, False, space_after=2, alignment=WD_ALIGN_PARAGRAPH.LEFT)

doc.add_paragraph()

# 4.2.3 State Chart Diagram
add_normal_text("4.2.3 State Chart Diagram", 12, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text("States:", 12, True, space_after=4, alignment=WD_ALIGN_PARAGRAPH.LEFT)
states = [
    "Idle: System waiting for input",
    "Sample Received: Behavior data loaded",
    "Semantic Processing: Converting features to natural language",
    "Binary Detection: Running GPT-2 binary classifier",
    "Family Classification: Running 12-class classifier (if ransomware)",
    "Result Display: Showing detection results and recommended actions",
]
for s in states:
    p = doc.add_paragraph(s, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

doc.add_paragraph()

# 4.2.4 System Architecture (Deployment Diagram equivalent)
add_normal_text("4.2.4 Deployment / System Architecture", 12, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "[Refer to Architecture Diagram in Chapter 1, Section 1.5 for the complete system architecture "
    "including the SRDC Shield Dual-Layer Client-Server framework.]",
    12, False, space_after=12
)

# 4.3 TECHNOLOGIES USED
add_normal_text("4.3 TECHNOLOGIES USED", 13, True, space_after=8, alignment=WD_ALIGN_PARAGRAPH.LEFT)

techs = [
    ("4.3.1 Python", "Python is the primary programming language used for the entire project. It provides extensive libraries for machine learning, data processing, and natural language processing. Python's simplicity and rich ecosystem make it the standard choice for AI/ML research implementations."),
    ("4.3.2 PyTorch", "PyTorch is an open-source deep learning framework developed by Meta AI. It provides tensor computation with strong GPU acceleration and a dynamic computational graph system. In this project, PyTorch is used to define the Classifier neural network, load pre-trained GPT-2 weights, fine-tune the model, and perform inference during the demo."),
    ("4.3.3 HuggingFace Transformers", "The Transformers library by HuggingFace provides access to thousands of pre-trained models. In this project, we use it to load the zhouce/RDC-GPT model - a GPT-2 variant that has been pre-trained on cybersecurity domain knowledge. The library handles tokenization, model architecture, and weight loading seamlessly."),
    ("4.3.4 GPT-2 Architecture", "GPT-2 (Generative Pre-trained Transformer 2) is a transformer-based language model developed by OpenAI. It uses self-attention mechanisms to understand contextual relationships in text. The model has 768-dimensional hidden states and processes input sequences up to 1024 tokens. In SRDC, we use GPT-2 not for text generation but as a feature extractor - the last hidden states are mean-pooled and passed through a linear classification head."),
    ("4.3.5 Google Colab", "Google Colaboratory provides free cloud-based GPU computing (NVIDIA T4). Since fine-tuning a GPT-2 model requires significant computational resources, all model training was performed on Google Colab. The trained model weights were then exported as .pth files for local inference."),
    ("4.3.6 Pandas & scikit-learn", "Pandas is used for reading, manipulating, and splitting the CSV dataset. scikit-learn provides the train_test_split function for creating stratified data splits, and classification_report and accuracy_score for evaluating model performance with precision, recall, and F1-score metrics."),
]
for title, desc in techs:
    add_normal_text(title, 12, True, space_after=4, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_normal_text(desc, 12, False, space_after=10)

add_page_break()

# ============================================================
# CHAPTER 5: IMPLEMENTATION
# ============================================================
add_centered_text("CHAPTER – 5", 16, True, space_after=6)
add_centered_text("IMPLEMENTATION", 16, True, space_after=18)

# 5.1
add_normal_text("5.1 Step 1: Environment Setup", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "A Python virtual environment was created to isolate project dependencies. PyTorch (CPU version), "
    "Transformers, Pandas, NumPy, and scikit-learn were installed. The official SRDC GitHub repository "
    "(Michael-zhouce/SRDC) was cloned to obtain the research scripts.",
    12, False, space_after=8
)

# 5.2
add_normal_text("5.2 Step 2: Dataset Acquisition", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "The Sgandurra et al. (2016) dataset was downloaded containing dynamic behavior logs of 582 ransomware "
    "samples across 11 families and 942 goodware samples (1,524 total). The VariableNames.txt file maps column "
    "indices to feature names (API names, registry keys, file extensions, etc.).",
    12, False, space_after=8
)

# 5.3
add_normal_text("5.3 Step 3: Internal Semantic Processing", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "The Internal_Semantic_Processing.py script was run to convert the raw CSV (binary 0/1 feature vectors) "
    "into natural language sentences. For example, API:GetFileSize=1 becomes 'get file size', and "
    "FILES:OPENED:C:\\file.doc becomes 'open the file C:\\file.doc'. This produced 7 semantic feature columns: "
    "apiFeatures, dropFeatures, regFeatures, filesFeatures, filesEXTFeatures, dirFeatures, and strFeatures.",
    12, False, space_after=4
)
add_normal_text("Output: after_feature_internal_semantic_process_data.csv (1,524 rows, ~10 MB)", 12, False, space_after=8)

# 5.4
add_normal_text("5.4 Step 4: Data Splitting", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "Two separate data splits were created for the two experiments described in the research paper:",
    12, False, space_after=4
)
add_normal_text(
    "Experiment 1 - Zero-Day Detection: 4 ransomware families were deliberately held out from training, "
    "appearing only in the test set. This simulates real-world zero-day threats where the model encounters "
    "families it has never seen. Files: zero_day_train.csv (1,388 rows) and zero_day_test.csv (268 rows).",
    12, False, space_after=4
)
add_normal_text(
    "Experiment 2 - Family Classification: A standard 80/20 stratified split keeping all 12 families in "
    "both sets. Files: train.csv (1,219 rows) and test.csv (305 rows).",
    12, False, space_after=8
)

# 5.5
add_normal_text("5.5 Step 5: Model Training on Google Colab (T4 GPU)", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "Both models were trained on Google Colab using a T4 GPU for 20 epochs each with Adam optimizer "
    "(learning rate: 1e-5), batch size 1, and CrossEntropyLoss.",
    12, False, space_after=4
)
add_normal_text(
    "Binary Detection Results (Zero-Day): The model achieved peak test accuracy of 97.01% across multiple "
    "epochs, successfully detecting ransomware from families it had never seen during training. The best "
    "model was saved as srdc_zero_day_BEST.pth.",
    12, False, space_after=4
)
add_normal_text(
    "Family Classification Results: The model achieved 97.29% training accuracy and 60.73% balanced accuracy "
    "at epoch 16. The lower balanced accuracy is expected due to severe class imbalance across the 12 families. "
    "The best model was saved as srdc_family_BEST.pth.",
    12, False, space_after=8
)

# 5.6
add_normal_text("5.6 Step 6: Interactive Demo", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "An interactive sandbox simulation (srdc_demo_fixed.py) was built that: (1) Loads both trained models, "
    "(2) Randomly selects 3 test samples (2 ransomware + 1 goodware), (3) Displays captured behavior "
    "(API calls, registry activity, file extensions), (4) Runs semantic analysis through the GPT-2 pipeline, "
    "and (5) Shows real-time predictions with confidence scores, true labels, and correctness.",
    12, False, space_after=8
)

# 5.7 SRDC Shield
add_normal_text("5.7 SRDC Shield – Production Browser Extension", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "To transition this research into a deployable cybersecurity application, the core AI engine was "
    'productized into a client-server system named "SRDC Shield". This architecture combines the fine-tuned '
    "deep learning models with a real-time Google Chrome Manifest V3 extension frontend and a local Python "
    "Flask backend service. The combined application actively intercepts high-risk web files before they "
    "execute, proving that semantic LLM models can move beyond offline validation to provide proactive "
    "perimeter protection.",
    12, False, space_after=8
)

# 5.8 Problems Faced & Solutions
add_normal_text("5.8 Problems Faced & Solutions", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)

problems = [
    ("1. UnicodeDecodeError:", "The VariableNames.txt file had special characters that Windows cp1252 encoding could not handle. Solution: Changed file read encoding to latin1."),
    ("2. Missing Dependencies:", "matplotlib was not initially installed. Solution: pip install matplotlib."),
    ("3. PowerShell Syntax Error:", "Multi-line commands with ^ continuation (a Bash syntax) failed in PowerShell. Solution: Used single-line commands and moved training to Google Colab."),
    ("4. No Local GPU:", "Training the GPT-2 model locally was infeasible without a GPU. Solution: Used Google Colab with T4 GPU."),
    ("5. Data Leakage:", "Initial data splits had zero-day families leaking into the training set. Solution: Wrote check_leakage.py to diagnose, then recreated proper leak-free splits in Colab."),
    ("6. Large Model Files:", "The .pth files (~498 MB each) exceeded GitHub's file size limit. Solution: Added *.pth to .gitignore."),
    ("7. Browser-Side Service Worker Crashes:", "Attempting to pull complex string data from large binary objects locked up the host CPU for up to 41 seconds, causing Chrome service worker timeouts. Solution: Implemented a structural 256KB Read Optimization, restricting feature harvesting to the initial 256 KB of PE file data, reducing latency from 41 seconds to under 0.05 seconds with zero loss in detection performance."),
    ("8. Archive Extraction Risks:", 'Attackers frequently drop malicious payloads inside nested archive patterns (.zip). Solution: Created an isolated "sandbox_temp/" workspace folder for secure archive extraction with automated cleanup after inference.'),
    ("9. Visual System Verification:", "Demonstrating ransomware containment safely during evaluations. Solution: Built an explicit 'simulate_zero_day' query handler that injects mock ransomware features for safe live demonstration."),
]
for title, desc in problems:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    run_t = p.add_run(title + " ")
    run_t.bold = True
    run_t.font.size = Pt(11)
    run_t.font.name = 'Times New Roman'
    run_d = p.add_run(desc)
    run_d.font.size = Pt(11)
    run_d.font.name = 'Times New Roman'

doc.add_paragraph()

# 5.9 Screenshots
add_normal_text("5.9 Screenshots", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text("[Insert UI Screenshots here]", 12, False, space_after=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_page_break()

# ============================================================
# CHAPTER 6: SOFTWARE TESTING
# ============================================================
add_centered_text("CHAPTER – 6", 16, True, space_after=6)
add_centered_text("SOFTWARE TESTING", 16, True, space_after=18)

# 6.1 Introduction
add_normal_text("6.1 Introduction", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "Software testing is a critical phase in the software development lifecycle that ensures the system "
    "works correctly and meets the research paper's expected performance benchmarks. Testing validates "
    "that the implementation faithfully reproduces the results claimed in the AAAI 2025 paper.",
    12, False, space_after=8
)

# 6.2 Testing Objectives
add_normal_text("6.2 Testing Objectives", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
objectives = [
    "Validate that semantic processing correctly converts binary features to natural language",
    "Ensure zero-day detection accuracy matches the paper's claims (97%)",
    "Verify family classification produces meaningful predictions (60.7% balanced accuracy)",
    "Confirm the demo runs end-to-end without errors",
    "Validate browser extension interception and scanning pipeline",
]
for obj in objectives:
    p = doc.add_paragraph(obj, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

doc.add_paragraph()

# 6.3 Testing Strategies
add_normal_text("6.3 Testing Strategies", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "Unit Testing: Individual components (semantic processing, tokenization, model inference) were tested "
    "in isolation to ensure correctness.",
    12, False, space_after=4
)
add_normal_text(
    "Integration Testing: The complete pipeline from raw data to prediction output was tested end-to-end "
    "to verify correct data flow between modules.",
    12, False, space_after=4
)
add_normal_text(
    "Performance Testing: Model inference time, feature extraction latency, and extension response times "
    "were measured to ensure real-time performance requirements are met.",
    12, False, space_after=8
)

# 6.4 System Evaluation
add_normal_text("6.4 System Evaluation", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
add_normal_text(
    "The system was evaluated on two metrics: (1) Binary detection accuracy of 97% on zero-day samples, "
    "and (2) Family classification balanced accuracy of 60.7% across 12 imbalanced classes. Both results "
    "align with the performance reported in the original research paper.",
    12, False, space_after=8
)

# 6.5 Test Cases
add_normal_text("6.5 Test Cases", 13, True, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)

# Test case table
tc_table = doc.add_table(rows=4, cols=4, style='Table Grid')
tc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

tc_headers = ['Test Case', 'Input', 'Expected Output', 'Actual Output']
for i, h in enumerate(tc_headers):
    cell = tc_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
    set_cell_shading(cell, 'D9E2F3')

test_cases = [
    ("TC1: Goodware Detection", "Goodware sample with typical system API calls (GetSystemDirectory, WriteConsole)", "SYSTEM IS CLEAN - GOODWARE", "SYSTEM IS CLEAN - GOODWARE (Confidence: 99.5%) ✓"),
    ("TC2: Ransomware (Trojan-Ransom)", "Ransomware sample with suspicious API calls (FreeVirtualMemory, CreateThread, OpenKey)", "RANSOMWARE DETECTED", "RANSOMWARE DETECTED (100.0%), Family: Trojan-Ransom (97.6%) ✓"),
    ("TC3: Ransomware (Reveton)", "Ransomware sample with network and file system activity", "RANSOMWARE DETECTED", "RANSOMWARE DETECTED (100.0%), Family: Reveton (99.9%) ✓"),
]
for row_idx, (tc, inp, expected, actual) in enumerate(test_cases, 1):
    tc_table.rows[row_idx].cells[0].text = tc
    tc_table.rows[row_idx].cells[1].text = inp
    tc_table.rows[row_idx].cells[2].text = expected
    tc_table.rows[row_idx].cells[3].text = actual
    for j in range(4):
        for p in tc_table.rows[row_idx].cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'

add_page_break()

# ============================================================
# CONCLUSION
# ============================================================
add_centered_text("CONCLUSION", 16, True, space_after=18)

add_normal_text(
    "The SRDC system successfully demonstrates the application of Large Language Models in cybersecurity "
    "for ransomware detection and classification. By converting raw program behavior into natural language "
    "and leveraging GPT-2's semantic understanding capabilities, the system achieves 97% accuracy in "
    "detecting zero-day ransomware - threats that traditional signature-based systems would completely miss.",
    12, False, space_after=8
)
add_normal_text(
    "The project validates the core thesis of the AAAI 2025 research paper: that semantic understanding of "
    "program behavior is fundamentally more powerful than statistical pattern matching for malware detection. "
    "The interactive sandbox demo provides a tangible, real-time visualization of the system's capabilities.",
    12, False, space_after=8
)
add_normal_text(
    "Key achievements of this implementation include: successful reproduction of the research pipeline, "
    "identification and resolution of data leakage issues, training of both binary and family classification "
    "models, and development of a polished interactive demo.",
    12, False, space_after=8
)

add_page_break()

# ============================================================
# FUTURE ENHANCEMENTS
# ============================================================
add_centered_text("FUTURE ENHANCEMENTS", 16, True, space_after=18)

enhancements = [
    "Real-time Sandbox Integration: Connect the model with a live sandbox environment (e.g., Cuckoo Sandbox) to intercept and analyze API calls in real-time as executables run.",
    "Automated Threat Mitigation: Implement a network kill-switch that automatically isolates infected machines when ransomware is detected with high confidence (>95%).",
    "Model Upgrades: Replace GPT-2 with smaller, more efficient models (e.g., DistilBERT, quantized Llama-3) for faster inference on edge devices without cloud connectivity.",
    "Web Dashboard: Build a web-based UI for security analysts to upload samples, view detection history, and monitor threat trends.",
    "Extended Dataset: Incorporate newer ransomware families beyond the original 11 families to improve the model's coverage of the modern threat landscape.",
    "Cloud Deployment: Deploy the detection API as a cloud service that endpoint agents can query for real-time threat analysis.",
    "Live OS Kernel Driver Attachment: Expand the local Flask engine into a background Windows service using file system mini-filter drivers to inspect disk events globally rather than just browser downloads.",
    "Multi-Model Ensemble Voting: Connect the current GPT-2 classifier head with lightweight, quantized configurations of newer models (e.g., Llama-3-8B-Instruct via local GGUF execution) to optimize complex family categorization scores.",
]
for e in enhancements:
    p = doc.add_paragraph(e, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

add_page_break()

# ============================================================
# REFERENCES
# ============================================================
add_centered_text("REFERENCES", 16, True, space_after=18)

references = [
    '[1] Ce Zhou et al., "SRDC: Semantics-based Ransomware Detection and Classification with LLM-assisted Pre-training", AAAI 2025.',
    '[2] Sgandurra, D., Munoz-Gonzalez, L., Mohsen, R., & Lupu, E.C. (2016). "Automated Dynamic Analysis of Ransomware: Benefits, Limitations and Use for Detection". arXiv:1609.03020.',
    '[3] Radford, A., et al. (2019). "Language Models are Unsupervised Multitask Learners". OpenAI. (GPT-2 Paper)',
    '[4] HuggingFace Transformers Documentation - https://huggingface.co/docs/transformers',
    '[5] PyTorch Documentation - https://pytorch.org/docs/stable/',
    '[6] zhouce/RDC-GPT Model - https://huggingface.co/zhouce/RDC-GPT',
]
for ref in references:
    add_normal_text(ref, 12, False, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)

add_page_break()

# ============================================================
# BIBLIOGRAPHY
# ============================================================
add_centered_text("BIBLIOGRAPHY", 16, True, space_after=18)

bibliography = [
    '[1] Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press.',
    '[2] Vaswani, A., et al. (2017). "Attention Is All You Need". NeurIPS.',
    '[3] Microsoft Security Documentation - https://docs.microsoft.com/en-us/security/',
    '[4] MITRE ATT&CK Framework - https://attack.mitre.org/',
    '[5] Google Colab Documentation - https://colab.research.google.com/',
]
for bib in bibliography:
    add_normal_text(bib, 12, False, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)


# ============================================================
# SAVE THE DOCUMENT
# ============================================================
output_path = r'c:\Users\sree nilay\Downloads\DOMAIN-PRO-SRDC\DOMAIN-PRO-SRDC\SRDC_DOCUMENTATION_IQAC_FORMAT.docx'
doc.save(output_path)
print(f"Document saved successfully to: {output_path}")
print("Done!")
